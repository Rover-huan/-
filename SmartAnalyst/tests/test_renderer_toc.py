from __future__ import annotations

import base64
import builtins
import copy
import logging
import re
import zipfile
from dataclasses import dataclass

import pytest
from docx import Document

from src import node4_renderer


@dataclass(frozen=True)
class _FakeSettings:
    enable_auto_toc: bool
    auto_toc_backend: str


def _read_docx_member(docx_path, member_name: str) -> str:
    with zipfile.ZipFile(docx_path) as archive:
        return archive.read(member_name).decode("utf-8")


def _paragraph_xml_containing(document_xml: str, text: str) -> str:
    match = re.search(rf"<w:p\b(?:(?!</w:p>).)*<w:t>{re.escape(text)}</w:t>(?:(?!</w:p>).)*</w:p>", document_xml)
    assert match is not None
    return match.group(0)


def _sample_image_path(tmp_path):
    image_path = tmp_path / "chart.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
        )
    )
    return str(image_path)


def _sample_result(tmp_path, task_id: int, question: str):
    return {
        "task_id": task_id,
        "question_zh": question,
        "image_path": _sample_image_path(tmp_path),
        "analysis_text": "图表显示指标存在阶段性变化。",
        "code_snippet": "print('demo')",
        "prepare_code": "df_clean = df.copy()",
        "plot_code": "plt.plot(df['年份'], df['失业率'])",
        "exploration_output": "(10, 3)",
        "cleaning_summary": "统一字段名称并删除缺失值。",
        "problem_solution": "处理了缺失值和字段编码问题。",
        "reflection_hint": "后续可以补充更多年份。",
        "column_mapping": {"year": "年份", "rate": "失业率"},
    }


def _sample_report_with_chart_notes():
    return {
        "section_1_intro": {
            "title": "一、引言与数据清洗说明",
            "content": "本报告基于课程数据进行量化分析，先说明数据处理过程。",
        },
        "section_2_analysis": [
            {
                "sub_title": "2022年全国城镇调查失业率的月度波动特征",
                "content": (
                    "本图展示失业率的月度波动。"
                    "完整可复现代码已保留在同步生成的 Jupyter Notebook 文件中，"
                    "DOCX 正文只保留方法说明、图表和文字分析。"
                ),
            },
            {
                "sub_title": "（一）16-24岁与25-59岁人口失业率的年龄组差异",
                "content": "本图比较不同年龄组失业率差异。",
            },
            {
                "sub_title": "(一)2003-2022年高工资省份与全国平均工资的长期增长趋势",
                "content": "本图展示工资长期增长趋势。",
            },
        ],
        "section_3_mechanism": {
            "title": "三、综合机制分析",
            "content": "综合前文图表，指标差异可能与结构性因素有关。",
        },
        "section_4_reflection": [
            {
                "sub_title": "四、遇到的问题及解决方法",
                "content": (
                    "关键代码与处理过程说明：代码完成读取、清洗和绘图。"
                    "完整可复现代码和每个 cell 的运行输出已保留在同步生成的 Jupyter Notebook 文件中。"
                ),
            },
            {
                "sub_title": "五、总结与思考",
                "content": "本报告完成了描述性分析，但仍缺少更严格的检验。",
            },
        ],
    }


def test_inserted_toc_field_is_present_in_docx(tmp_path):
    docx_path = tmp_path / "toc.docx"
    document = Document()

    node4_renderer._configure_document_styles(document)
    node4_renderer._enable_update_fields_on_open(document)
    node4_renderer._insert_toc_page(document)
    node4_renderer._add_section_heading(
        document,
        "一、引言与数据清洗说明",
        level=1,
        start_on_new_page=False,
    )
    node4_renderer._add_section_heading(
        document,
        "（一）销售额变化趋势分析",
        level=2,
        start_on_new_page=False,
    )
    document.save(docx_path)

    document_xml = _read_docx_member(docx_path, "word/document.xml")

    assert "<w:t>目录</w:t>" in document_xml
    assert 'TOC \\o "1-2" \\h \\z \\u' in document_xml
    assert '<w:outlineLvl w:val="0"/>' in document_xml
    assert '<w:outlineLvl w:val="1"/>' in document_xml


def test_section_2_subheadings_are_numbered_for_body_and_toc(tmp_path):
    docx_path = tmp_path / "report.docx"
    Document().save(docx_path)
    results = [
        _sample_result(tmp_path, 1, "2022年失业率如何波动？"),
        _sample_result(tmp_path, 2, "不同年龄组失业率是否存在差异？"),
        _sample_result(tmp_path, 3, "工资是否存在长期增长趋势？"),
    ]
    report_data = _sample_report_with_chart_notes()
    original_report_data = copy.deepcopy(report_data)

    node4_renderer._rebuild_report_body(docx_path, results, report_data)

    document = Document(docx_path)
    paragraph_texts = [paragraph.text for paragraph in document.paragraphs]

    assert "一、引言与数据清洗说明" in paragraph_texts
    assert node4_renderer.ANALYSIS_SECTION_TITLE in paragraph_texts
    assert "三、综合机制分析" in paragraph_texts
    assert "四、遇到的问题及解决方法" in paragraph_texts
    assert "五、总结与思考" in paragraph_texts
    assert "（一）2022年全国城镇调查失业率的月度波动特征" in paragraph_texts
    assert "（二）16-24岁与25-59岁人口失业率的年龄组差异" in paragraph_texts
    assert "（三）2003-2022年高工资省份与全国平均工资的长期增长趋势" in paragraph_texts
    assert "（一）（一）16-24岁与25-59岁人口失业率的年龄组差异" not in paragraph_texts

    document_xml = _read_docx_member(docx_path, "word/document.xml")
    first_subheading_xml = _paragraph_xml_containing(
        document_xml,
        "（一）2022年全国城镇调查失业率的月度波动特征",
    )

    assert 'TOC \\o "1-2" \\h \\z \\u' in document_xml
    assert '<w:outlineLvl w:val="1"/>' in first_subheading_xml
    assert report_data == original_report_data


def test_chart_sections_drop_repeated_notebook_note_but_reflection_keeps_one(tmp_path):
    docx_path = tmp_path / "report.docx"
    Document().save(docx_path)
    results = [
        _sample_result(tmp_path, 1, "2022年失业率如何波动？"),
        _sample_result(tmp_path, 2, "不同年龄组失业率是否存在差异？"),
        _sample_result(tmp_path, 3, "工资是否存在长期增长趋势？"),
    ]

    node4_renderer._rebuild_report_body(docx_path, results, _sample_report_with_chart_notes())

    document = Document(docx_path)
    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    section_2_text = body_text.split("三、综合机制分析", 1)[0]
    section_4_text = body_text.split("四、遇到的问题及解决方法", 1)[1]

    assert "完整可复现代码已保留在同步生成的 Jupyter Notebook 文件中" not in section_2_text
    assert "DOCX 正文只保留方法说明、图表和文字分析" not in section_2_text
    assert "完整可复现代码和每个 cell 的运行输出已保留在同步生成的 Jupyter Notebook 文件中" in section_4_text
    assert body_text.count("完整可复现代码和每个 cell 的运行输出已保留在同步生成的 Jupyter Notebook 文件中") == 1


def test_existing_toc_field_gets_plain_centered_title(tmp_path):
    docx_path = tmp_path / "existing-toc-with-title.docx"
    document = Document()
    toc_field_paragraph = document.add_paragraph()
    node4_renderer._append_toc_field(toc_field_paragraph)

    inserted = node4_renderer._ensure_toc_title_before_existing_toc(document)

    assert inserted is True
    assert document.paragraphs[0].text == "目录"
    assert document.paragraphs[0].alignment == node4_renderer.WD_ALIGN_PARAGRAPH.CENTER
    assert document.paragraphs[0].style.name not in {"Heading 1", "Heading 2", "标题1", "标题2"}

    document.save(docx_path)
    document_xml = _read_docx_member(docx_path, "word/document.xml")
    title_xml = _paragraph_xml_containing(document_xml, "目录")

    assert 'TOC \\o "1-2" \\h \\z \\u' in document_xml
    assert "w:outlineLvl" not in title_xml


@pytest.mark.parametrize("title_text", ["目录", "目 录", "目　录"])
def test_existing_toc_title_is_not_inserted_twice(title_text):
    document = Document()
    document.add_paragraph(title_text)
    toc_field_paragraph = document.add_paragraph()
    node4_renderer._append_toc_field(toc_field_paragraph)

    inserted = node4_renderer._ensure_toc_title_before_existing_toc(document)

    assert inserted is False
    assert [paragraph.text for paragraph in document.paragraphs].count(title_text) == 1
    assert len(document.paragraphs) == 2


def test_update_fields_on_open_is_enabled_in_settings(tmp_path):
    docx_path = tmp_path / "update-fields.docx"
    document = Document()

    node4_renderer._enable_update_fields_on_open(document)
    document.save(docx_path)

    settings_xml = _read_docx_member(docx_path, "word/settings.xml")

    assert "<w:updateFields w:val=\"true\"" in settings_xml


def test_refresh_docx_toc_returns_false_when_win32com_is_unavailable(monkeypatch, tmp_path):
    docx_path = tmp_path / "report.docx"
    Document().save(docx_path)

    monkeypatch.setattr(node4_renderer.platform, "system", lambda: "Windows")
    monkeypatch.setattr(
        node4_renderer,
        "get_settings",
        lambda: _FakeSettings(enable_auto_toc=True, auto_toc_backend="word_com"),
    )
    real_import = builtins.__import__

    def fake_import(name, globals=None, locals=None, fromlist=(), level=0):
        if name.startswith("win32com") or name == "pythoncom":
            raise ImportError(name)
        return real_import(name, globals, locals, fromlist, level)

    monkeypatch.setattr(builtins, "__import__", fake_import)

    assert node4_renderer.refresh_docx_toc(docx_path) is False


def test_refresh_docx_toc_returns_false_without_calling_backend_when_config_disabled(
    monkeypatch,
    tmp_path,
):
    docx_path = tmp_path / "report.docx"
    Document().save(docx_path)

    monkeypatch.setattr(
        node4_renderer,
        "get_settings",
        lambda: _FakeSettings(enable_auto_toc=False, auto_toc_backend="word_com"),
    )

    def fail_if_backend_called(_file_path):
        raise AssertionError("TOC backend should not be called when auto TOC is disabled")

    monkeypatch.setattr(node4_renderer, "_refresh_toc_with_word_com", fail_if_backend_called)

    assert node4_renderer.refresh_docx_toc(docx_path) is False


def test_refresh_docx_toc_returns_false_without_calling_backend_when_backend_none(
    monkeypatch,
    tmp_path,
):
    docx_path = tmp_path / "report.docx"
    Document().save(docx_path)

    monkeypatch.setattr(
        node4_renderer,
        "get_settings",
        lambda: _FakeSettings(enable_auto_toc=True, auto_toc_backend="none"),
    )

    def fail_if_backend_called(_file_path):
        raise AssertionError("TOC backend should not be called for AUTO_TOC_BACKEND=none")

    monkeypatch.setattr(node4_renderer, "_refresh_toc_with_word_com", fail_if_backend_called)

    assert node4_renderer.refresh_docx_toc(docx_path) is False


def test_refresh_docx_toc_libreoffice_backend_is_reserved(monkeypatch, tmp_path, caplog):
    docx_path = tmp_path / "report.docx"
    Document().save(docx_path)

    monkeypatch.setattr(
        node4_renderer,
        "get_settings",
        lambda: _FakeSettings(enable_auto_toc=True, auto_toc_backend="libreoffice"),
    )

    with caplog.at_level(logging.WARNING):
        assert node4_renderer.refresh_docx_toc(docx_path) is False

    assert "libreoffice backend 尚未实现" in caplog.text


def test_refresh_docx_toc_aspose_backend_is_reserved(monkeypatch, tmp_path, caplog):
    docx_path = tmp_path / "report.docx"
    Document().save(docx_path)

    monkeypatch.setattr(
        node4_renderer,
        "get_settings",
        lambda: _FakeSettings(enable_auto_toc=True, auto_toc_backend="aspose"),
    )

    with caplog.at_level(logging.WARNING):
        assert node4_renderer.refresh_docx_toc(docx_path) is False

    assert "aspose backend 尚未实现" in caplog.text
