from __future__ import annotations

import base64
import zipfile

from docx import Document

from src import node4_renderer


def _read_docx_member(docx_path, member_name: str) -> str:
    with zipfile.ZipFile(docx_path) as archive:
        return archive.read(member_name).decode("utf-8")


def _paragraph_has_page_break(paragraph) -> bool:
    return 'w:type="page"' in paragraph._p.xml


def _find_paragraph(document: Document, text: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise AssertionError(f"Paragraph not found: {text}")


def _previous_paragraph(document: Document, paragraph):
    paragraphs = list(document.paragraphs)
    for index, candidate in enumerate(paragraphs):
        if candidate._p is paragraph._p:
            assert index > 0
            return paragraphs[index - 1]
    raise AssertionError(f"Paragraph not found in document: {paragraph.text}")


def _sample_results(tmp_path):
    image_path = tmp_path / "chart.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
        )
    )
    return [
        {
            "task_id": 1,
            "question_zh": "地区工资水平是否存在差异？",
            "image_path": str(image_path),
            "analysis_text": "图表显示不同地区工资水平存在差异。",
            "code_snippet": "print('demo')",
            "prepare_code": "df_clean = df.copy()",
            "plot_code": "plt.savefig(output_image_path, dpi=300, bbox_inches='tight')",
            "exploration_output": "(10, 3)",
            "cleaning_summary": "统一字段名称并删除缺失值。",
            "problem_solution": "处理了缺失值和字段编码问题。",
            "reflection_hint": "后续可以补充更多年份。",
            "column_mapping": {"wage": "工资"},
        }
    ]


def _sample_report():
    return {
        "section_1_intro": {
            "title": "一、引言与数据说明",
            "content": "本报告基于课程数据进行量化分析，先说明数据处理过程。",
        },
        "section_2_analysis": [
            {
                "sub_title": "（一）地区工资差异分析",
                "content": "从图中可以看出，不同地区工资水平呈现一定差异。",
            }
        ],
        "section_3_mechanism": {
            "title": "三、综合机制分析",
            "content": "综合前文图表，地区差异可能与产业结构和人口流动有关。",
        },
        "section_4_reflection": [
            {
                "sub_title": "四、遇到的问题及解决方法",
                "content": "数据处理中主要遇到字段编码、缺失值和地区名称不一致等问题。",
            },
            {
                "sub_title": "五、总结与思考",
                "content": "本报告完成了描述性分析，但仍缺少更严格的回归检验。",
            },
        ],
    }


def test_late_major_sections_do_not_start_with_hard_page_breaks(tmp_path):
    docx_path = tmp_path / "report.docx"
    Document().save(docx_path)

    node4_renderer._rebuild_report_body(docx_path, _sample_results(tmp_path), _sample_report())

    document = Document(docx_path)
    for heading_text in [
        node4_renderer.ANALYSIS_SECTION_TITLE,
        "三、综合机制分析",
        "四、遇到的问题及解决方法",
        "五、总结与思考",
    ]:
        heading = _find_paragraph(document, heading_text)
        assert heading.paragraph_format.page_break_before is not True
        assert "w:pageBreakBefore" not in heading._p.xml
        assert not _paragraph_has_page_break(_previous_paragraph(document, heading))


def test_toc_page_breaks_and_title_survive_natural_body_layout(tmp_path):
    docx_path = tmp_path / "report.docx"
    Document().save(docx_path)

    node4_renderer._rebuild_report_body(docx_path, _sample_results(tmp_path), _sample_report())

    document_xml = _read_docx_member(docx_path, "word/document.xml")

    assert "<w:t>目录</w:t>" in document_xml
    assert 'TOC \\o "1-2" \\h \\z \\u' in document_xml
    assert document_xml.count('w:type="page"') == 2


def test_major_headings_keep_with_next_without_forcing_new_pages(tmp_path):
    docx_path = tmp_path / "report.docx"
    Document().save(docx_path)

    node4_renderer._rebuild_report_body(docx_path, _sample_results(tmp_path), _sample_report())

    document = Document(docx_path)
    heading = _find_paragraph(document, "三、综合机制分析")

    assert heading.paragraph_format.keep_with_next is True
    assert heading.paragraph_format.page_break_before is not True
    assert "w:pageBreakBefore" not in heading._p.xml
