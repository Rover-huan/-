"""Node 4 renderer for template-driven report, notebook, and cleaning-summary outputs."""

from __future__ import annotations

import base64
import json
import logging
import os
import platform
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, TypedDict

NBF_IMPORT_ERROR: ImportError | None = None
DOCX_TEMPLATE_IMPORT_ERROR: ImportError | None = None
DOCX_IMPORT_ERROR: ImportError | None = None

try:
    import nbformat as nbf
except ImportError as exc:
    NBF_IMPORT_ERROR = exc
    nbf = None  # type: ignore[assignment]

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt
except ImportError as exc:
    DOCX_IMPORT_ERROR = exc
    Document = None  # type: ignore[assignment]
    WD_STYLE_TYPE = None  # type: ignore[assignment]
    WD_ALIGN_PARAGRAPH = None  # type: ignore[assignment]
    WD_BREAK = None  # type: ignore[assignment]
    OxmlElement = None  # type: ignore[assignment]
    qn = None  # type: ignore[assignment]
    Mm = None  # type: ignore[assignment]
    Pt = None  # type: ignore[assignment]

try:
    from docxtpl import DocxTemplate, InlineImage
except ImportError as exc:
    DOCX_TEMPLATE_IMPORT_ERROR = exc
    DocxTemplate = None  # type: ignore[assignment]
    InlineImage = None  # type: ignore[assignment]

try:
    from docx2pdf import convert
except ImportError:
    convert = None  # type: ignore[assignment]

try:
    from service.config import get_settings
except ImportError:
    get_settings = None  # type: ignore[assignment]


logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logging.getLogger().setLevel(logging.INFO)
LOGGER = logging.getLogger(__name__)

MODULE_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = MODULE_DIR.parent
TEMPLATE_PATH = PROJECT_ROOT / "templates" / "report_template.docx"
OUTPUTS_DIR = PROJECT_ROOT / "outputs"

BODY_FONT_NAME = "SimSun"
BODY_EAST_ASIA_FONT = "宋体"
BODY_FONT_SIZE_PT = 12
BODY_LINE_SPACING = 1.5
TOC_TITLE_TEXT = "目 录"
TOC_FIELD_INSTRUCTION = 'TOC \\o "1-2" \\h \\z \\u'
MAJOR_SECTION_PATTERN = re.compile(r"^[一二三四五六七八九十]+、")
TASK_SECTION_PATTERN = re.compile(r"^任务\s*\d+\s*[：:]")
ANALYSIS_SECTION_TITLE = "二、量化分析与可视化"
CODE_FONT_NAME = "Consolas"
CODE_FONT_SIZE_PT = 10
CODE_BACKGROUND_FILL = "F3F3F3"
CODE_BORDER_COLOR = "D9D9D9"

REQUIRED_RESULT_KEYS = {
    "task_id",
    "question_zh",
    "image_path",
    "analysis_text",
    "code_snippet",
    "prepare_code",
    "plot_code",
    "exploration_output",
    "cleaning_summary",
    "problem_solution",
    "reflection_hint",
    "column_mapping",
}
REQUIRED_REPORT_KEYS = {
    "section_1_intro",
    "section_2_analysis",
    "section_3_mechanism",
    "section_4_reflection",
}
REQUIRED_DATA_SUMMARY_KEYS = {
    "dataset_path",
    "dataset_name",
    "file_type",
    "shape_text",
    "info_text",
    "missing_summary_text",
    "duplicate_count_text",
    "preview_text",
    "load_code",
}


class ExecutionResult(TypedDict):
    """One rendered analysis result from Node 3 plus renderer enrichment."""

    task_id: int
    question_zh: str
    image_path: str
    analysis_text: str
    code_snippet: str
    prepare_code: str
    plot_code: str
    exploration_output: str
    cleaning_summary: str
    problem_solution: str
    reflection_hint: str
    column_mapping: dict[str, str]


class ReportSectionBlock(TypedDict):
    """One major report section with title and content."""

    title: str
    content: str


class ReportSectionItem(TypedDict):
    """One report subsection item."""

    sub_title: str
    content: str


class ReportData(TypedDict):
    """TOC-friendly report data from Node 3.5."""

    section_1_intro: ReportSectionBlock
    section_2_analysis: list[ReportSectionItem]
    section_3_mechanism: ReportSectionBlock
    section_4_reflection: list[ReportSectionItem]


class DataSummary(TypedDict):
    """Dataset summary payload created in main.py for renderer reuse."""

    dataset_path: str
    dataset_name: str
    file_type: str
    shape_text: str
    info_text: str
    missing_summary_text: str
    duplicate_count_text: str
    preview_text: str
    load_code: str


class RenderArtifacts(TypedDict):
    """Paths to generated report artifacts."""

    docx_path: str
    pdf_path: str | None
    notebook_path: str
    cleaning_summary_path: str


class TemplateLoopItem(TypedDict):
    """One loop item consumed by the final Word template."""

    question_zh: str
    analysis_text: str
    source_code: str
    image: Any


@dataclass(frozen=True)
class RendererPaths:
    """Concrete output paths for one render invocation."""

    output_dir: Path
    docx_path: Path
    pdf_path: Path
    notebook_path: Path
    cleaning_summary_path: Path


def _build_renderer_paths(output_dir: str | Path | None = None) -> RendererPaths:
    """Resolve per-job renderer output paths."""
    resolved_output_dir = Path(output_dir or OUTPUTS_DIR).resolve()
    return RendererPaths(
        output_dir=resolved_output_dir,
        docx_path=resolved_output_dir / "Final_Report.docx",
        pdf_path=resolved_output_dir / "Final_Report.pdf",
        notebook_path=resolved_output_dir / "Analysis_Notebook.ipynb",
        cleaning_summary_path=resolved_output_dir / "Cleaning_Steps_Summary.txt",
    )


def _require_docx_runtime() -> None:
    """Ensure python-docx level dependencies are available."""
    if (
        Document is None
        or WD_STYLE_TYPE is None
        or WD_ALIGN_PARAGRAPH is None
        or WD_BREAK is None
        or OxmlElement is None
        or qn is None
        or Mm is None
        or Pt is None
    ):
        raise RuntimeError(
            "python-docx is not installed, so Word post-processing is unavailable."
        ) from DOCX_IMPORT_ERROR


def _new_notebook() -> Any:
    """Create a notebook object with nbformat when available, else a JSON fallback."""
    if nbf is not None:
        return nbf.v4.new_notebook()
    return {
        "cells": [],
        "metadata": {},
        "nbformat": 4,
        "nbformat_minor": 5,
    }


def _new_markdown_cell(source: str) -> Any:
    """Create a markdown cell with nbformat when available, else a JSON fallback."""
    if nbf is not None:
        return nbf.v4.new_markdown_cell(source)
    return {
        "cell_type": "markdown",
        "metadata": {},
        "source": source,
    }


def _new_code_cell(source: str) -> Any:
    """Create a code cell with nbformat when available, else a JSON fallback."""
    if nbf is not None:
        return nbf.v4.new_code_cell(source=source)
    return {
        "cell_type": "code",
        "execution_count": None,
        "metadata": {},
        "outputs": [],
        "source": source,
    }


def _new_output(output_type: str, **payload: Any) -> Any:
    """Create a notebook output record."""
    if nbf is not None:
        return nbf.v4.new_output(output_type=output_type, **payload)
    return {
        "output_type": output_type,
        **payload,
    }


def _write_notebook(notebook: Any, notebook_file: Any) -> None:
    """Write a notebook using nbformat when available, else raw JSON."""
    if nbf is not None:
        nbf.write(notebook, notebook_file)
        return
    json.dump(notebook, notebook_file, ensure_ascii=False, indent=1)
    notebook_file.write("\n")


def _ensure_outputs_dir(output_dir: Path) -> None:
    """Create the outputs directory if it does not already exist."""
    output_dir.mkdir(parents=True, exist_ok=True)


def _set_run_font(
    run: Any,
    *,
    font_name: str = BODY_FONT_NAME,
    east_asia_font: str = BODY_EAST_ASIA_FONT,
    size_pt: int = BODY_FONT_SIZE_PT,
    bold: bool | None = None,
) -> None:
    """Apply consistent run-level font settings."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    run_properties = run._r.get_or_add_rPr()
    run_fonts = run_properties.rFonts
    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        run_properties.append(run_fonts)
    run_fonts.set(qn("w:ascii"), font_name)
    run_fonts.set(qn("w:hAnsi"), font_name)
    run_fonts.set(qn("w:eastAsia"), east_asia_font)


def _get_style_by_name(document: Any, style_name: str) -> Any | None:
    """Return a paragraph style by name when it exists."""
    try:
        return document.styles[style_name]
    except KeyError:
        return None


def _set_style_font(
    style: Any,
    *,
    font_name: str,
    east_asia_font: str,
    size_pt: int,
    bold: bool = False,
    line_spacing: float | None = None,
) -> None:
    """Apply font settings to a paragraph style."""
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    run_properties = style._element.get_or_add_rPr()
    run_fonts = run_properties.rFonts
    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        run_properties.append(run_fonts)
    run_fonts.set(qn("w:ascii"), font_name)
    run_fonts.set(qn("w:hAnsi"), font_name)
    run_fonts.set(qn("w:eastAsia"), east_asia_font)
    if line_spacing is not None:
        style.paragraph_format.line_spacing = line_spacing


def _set_style_outline_level(style: Any, outline_level: int) -> None:
    """Set Word outline level metadata so TOC fields can find headings."""
    paragraph_properties = style._element.get_or_add_pPr()
    outline = paragraph_properties.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        paragraph_properties.append(outline)
    outline.set(qn("w:val"), str(outline_level))


def _set_paragraph_outline_level(paragraph: Any, outline_level: int) -> None:
    """Set paragraph-level outline metadata for robust TOC generation."""
    paragraph_properties = paragraph._p.get_or_add_pPr()
    outline = paragraph_properties.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        paragraph_properties.append(outline)
    outline.set(qn("w:val"), str(outline_level))


def _ensure_heading_styles(document: Any) -> tuple[Any, Any]:
    """Ensure standard Heading 1 and Heading 2 styles exist for TOC generation."""
    heading_1_style = _get_style_by_name(document, "Heading 1")
    heading_2_style = _get_style_by_name(document, "Heading 2")

    if heading_1_style is None:
        template_heading_1 = _get_style_by_name(document, "标题1")
        if template_heading_1 is not None:
            template_heading_1.name = "heading 1"
            heading_1_style = template_heading_1
        else:
            heading_1_style = document.styles.add_style("heading 1", WD_STYLE_TYPE.PARAGRAPH)

    if heading_2_style is None:
        template_heading_2 = _get_style_by_name(document, "标题2")
        if template_heading_2 is not None:
            template_heading_2.name = "heading 2"
            heading_2_style = template_heading_2
        else:
            heading_2_style = document.styles.add_style("heading 2", WD_STYLE_TYPE.PARAGRAPH)

    _set_style_font(
        heading_1_style,
        font_name=BODY_FONT_NAME,
        east_asia_font=BODY_EAST_ASIA_FONT,
        size_pt=16,
        bold=True,
        line_spacing=1.15,
    )
    _set_style_outline_level(heading_1_style, 0)
    _set_style_font(
        heading_2_style,
        font_name=BODY_FONT_NAME,
        east_asia_font=BODY_EAST_ASIA_FONT,
        size_pt=14,
        bold=True,
        line_spacing=1.15,
    )
    _set_style_outline_level(heading_2_style, 1)
    return heading_1_style, heading_2_style


def _configure_document_styles(document: Any) -> None:
    """Configure document styles for body text and standard headings."""
    normal_style = document.styles["Normal"]
    _set_style_font(
        normal_style,
        font_name=BODY_FONT_NAME,
        east_asia_font=BODY_EAST_ASIA_FONT,
        size_pt=BODY_FONT_SIZE_PT,
        bold=False,
        line_spacing=BODY_LINE_SPACING,
    )
    _ensure_heading_styles(document)


def _is_template_body_start(paragraph: Any) -> bool:
    """Return True when the paragraph marks the start of the templated body section."""
    style_name = getattr(paragraph.style, "name", "")
    text = paragraph.text.strip()
    return style_name in {"Heading 1", "标题1"} or bool(MAJOR_SECTION_PATTERN.match(text))


def _remove_paragraph(paragraph: Any) -> None:
    """Delete a paragraph from the document."""
    paragraph._element.getparent().remove(paragraph._element)


def _clear_template_body(document: Any) -> None:
    """Keep the cover page from the template and remove the pre-authored body."""
    removing = False
    for paragraph in list(document.paragraphs):
        if not removing and _is_template_body_start(paragraph):
            removing = True
        if removing:
            _remove_paragraph(paragraph)


def _append_toc_field(paragraph: Any) -> None:
    """Insert a Table of Contents field into a paragraph."""
    begin_run = paragraph.add_run()
    begin_field = OxmlElement("w:fldChar")
    begin_field.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin_field)

    instruction_run = paragraph.add_run()
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = TOC_FIELD_INSTRUCTION
    instruction_run._r.append(instruction)

    separate_run = paragraph.add_run()
    separate_field = OxmlElement("w:fldChar")
    separate_field.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate_field)

    placeholder_run = paragraph.add_run("按 F9 更新目录")
    _set_run_font(placeholder_run)

    end_field = OxmlElement("w:fldChar")
    end_field.set(qn("w:fldCharType"), "end")
    placeholder_run._r.append(end_field)


def _enable_update_fields_on_open(document: Any) -> None:
    """Ask Word to update document fields on open as a fallback."""
    settings_element = document.settings._element
    update_fields = settings_element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings_element.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _normalize_existing_toc_fields(document: Any) -> int:
    """Normalize existing template TOC field instructions and return their count."""
    toc_count = 0
    for paragraph in document.paragraphs:
        for instruction in paragraph._p.iter(qn("w:instrText")):
            if instruction.text and "TOC" in instruction.text:
                instruction.text = f" {TOC_FIELD_INSTRUCTION} "
                toc_count += 1
        for simple_field in paragraph._p.iter(qn("w:fldSimple")):
            instruction_text = simple_field.get(qn("w:instr"))
            if instruction_text and "TOC" in instruction_text:
                simple_field.set(qn("w:instr"), TOC_FIELD_INSTRUCTION)
                toc_count += 1
    return toc_count


def _insert_toc_page(document: Any) -> None:
    """Insert a TOC page after the cover."""
    document.add_page_break()
    toc_title_paragraph = document.add_paragraph()
    toc_title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title_run = toc_title_paragraph.add_run(TOC_TITLE_TEXT)
    _set_run_font(toc_title_run, size_pt=14, bold=True)
    toc_title_paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING

    toc_field_paragraph = document.add_paragraph()
    toc_field_paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
    _append_toc_field(toc_field_paragraph)

    document.add_page_break()


def _apply_body_paragraph_format(paragraph: Any, *, code_mode: bool = False) -> None:
    """Apply academic body text formatting to a paragraph."""
    paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        if code_mode:
            _set_run_font(
                run,
                font_name=CODE_FONT_NAME,
                east_asia_font=CODE_FONT_NAME,
                size_pt=CODE_FONT_SIZE_PT,
                bold=False,
            )
        else:
            _set_run_font(run)


def _add_body_paragraphs(document: Any, text: str) -> None:
    """Render plain body text while preserving paragraph breaks."""
    for raw_paragraph in text.splitlines():
        paragraph_text = raw_paragraph.strip()
        if not paragraph_text:
            continue
        paragraph = document.add_paragraph(paragraph_text)
        _apply_body_paragraph_format(paragraph)


def _add_label_paragraph(document: Any, label_text: str) -> None:
    """Add a bold body label paragraph."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(label_text)
    _set_run_font(run, bold=True)
    _apply_body_paragraph_format(paragraph)


def _set_cell_shading(cell: Any, fill_color: str) -> None:
    """Apply background shading to a table cell."""
    tc_properties = cell._tc.get_or_add_tcPr()
    shading = tc_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_properties.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill_color)


def _set_cell_borders(cell: Any, border_color: str) -> None:
    """Apply a subtle border to a table cell."""
    tc_properties = cell._tc.get_or_add_tcPr()
    borders = tc_properties.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_properties.append(borders)

    for edge_name in ("top", "left", "bottom", "right"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "6")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), border_color)


def _add_code_block(document: Any, source_code: str) -> None:
    """Render source code inside a shaded one-cell table."""
    table = document.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    _set_cell_shading(cell, CODE_BACKGROUND_FILL)
    _set_cell_borders(cell, CODE_BORDER_COLOR)

    paragraph = cell.paragraphs[0]
    paragraph.text = ""
    paragraph.paragraph_format.line_spacing = 1.1
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(source_code.replace("\t", "    "))
    _set_run_font(
        run,
        font_name=CODE_FONT_NAME,
        east_asia_font=CODE_FONT_NAME,
        size_pt=CODE_FONT_SIZE_PT,
        bold=False,
    )

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def _add_section_heading(document: Any, heading_text: str, *, level: int, start_on_new_page: bool) -> Any:
    """Add a standard Word heading and optionally start it on a fresh page."""
    if start_on_new_page:
        document.add_page_break()
    heading = document.add_heading(heading_text, level=level)
    _set_paragraph_outline_level(heading, 0 if level == 1 else 1)
    for run in heading.runs:
        _set_run_font(
            run,
            font_name=BODY_FONT_NAME,
            east_asia_font=BODY_EAST_ASIA_FONT,
            size_pt=16 if level == 1 else 14,
            bold=True,
        )
    return heading


def _add_figure(document: Any, image_path: str) -> None:
    """Insert a centered figure into the report."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(image_path, width=Mm(145))
    paragraph.paragraph_format.space_after = Pt(6)


def _is_non_core_method_field(field_name: str) -> bool:
    """Return True for helper columns that should not appear in the method summary."""
    normalized = field_name.strip().lower()
    if not normalized:
        return True
    helper_tokens = (
        "英文",
        "代码",
        "编码",
        "编号",
        "序号",
        "index",
        "id",
        "code",
        "english",
    )
    return any(token in normalized for token in helper_tokens)


def _select_method_fields(
    result: ExecutionResult,
    *,
    max_fields: int = 3,
    prefer_question: bool = True,
) -> list[str]:
    """Pick a short list of fields that are likely central to the rendered chart."""
    field_pairs = [
        (str(source_name).strip(), str(target_name).strip())
        for source_name, target_name in result["column_mapping"].items()
        if str(source_name).strip()
        and str(target_name).strip()
        and not _is_non_core_method_field(str(target_name))
    ]
    question_text = result["question_zh"]
    code_text = "\n".join([result["plot_code"], result["prepare_code"], result["analysis_text"]])

    selected: list[str] = []

    def add_field(field_name: str) -> None:
        if field_name and field_name not in selected:
            selected.append(field_name)

    if prefer_question:
        for source_name, target_name in field_pairs:
            if target_name in question_text or source_name in question_text:
                add_field(target_name)

    quoted_names = re.findall(r"['\"]([^'\"]{1,40})['\"]", code_text)
    for quoted_name in quoted_names:
        for source_name, target_name in field_pairs:
            if quoted_name in {source_name, target_name}:
                add_field(target_name)
                break
        if len(selected) >= max_fields:
            return selected[:max_fields]

    for _source_name, target_name in field_pairs:
        add_field(target_name)
        if len(selected) >= max_fields:
            break

    return selected[:max_fields]


def _looks_like_time_field(field_name: str) -> bool:
    """Return True for fields that usually represent time or order."""
    return any(token in field_name for token in ("时间", "日期", "年份", "年度", "月份", "季度", "周", "序号"))


def _build_analysis_method_text(result: ExecutionResult) -> str:
    """Summarize the chart method for DOCX without embedding full Python code."""
    question = result["question_zh"]
    code_text = "\n".join([result["plot_code"], result["analysis_text"]]).lower()
    code_note = "完整可复现代码已保留在同步生成的 Jupyter Notebook 文件中，DOCX 正文只保留方法说明、图表和文字分析。"

    if "scatter" in code_text:
        field_names = _select_method_fields(result, max_fields=2, prefer_question=False)
        if len(field_names) >= 2:
            return (
                f"本图以 {field_names[0]} 为横轴、{field_names[1]} 为纵轴，"
                "通过散点分布观察二者是否存在同步变化、离群点或区域差异。"
                f"{code_note}"
            )

    if "bar(" in code_text or "barh(" in code_text or "groupby" in code_text:
        field_names = _select_method_fields(result, max_fields=2, prefer_question=False)
        if len(field_names) >= 2:
            return (
                f"本图以 {field_names[0]} 为分组维度，对 {field_names[1]} 进行比较，"
                "用于观察不同地区或类别之间的差异。"
                f"{code_note}"
            )

    if "plot(" in code_text or "line" in code_text:
        field_names = _select_method_fields(result, max_fields=3, prefer_question=False)
        time_fields = [field_name for field_name in field_names if _looks_like_time_field(field_name)]
        if time_fields:
            metric_fields = [field_name for field_name in field_names if field_name not in time_fields]
            metric_text = metric_fields[0] if metric_fields else "核心指标"
            return (
                f"本图按 {time_fields[0]} 展示 {metric_text} 的变化，"
                "用于观察该指标是否存在阶段性波动或趋势。"
                f"{code_note}"
            )

    field_names = _select_method_fields(result)
    if len(field_names) >= 2:
        return (
            f"本图围绕 {field_names[0]} 与 {field_names[1]} 的关系展开分析，"
            "在完成必要的数据清洗后进行可视化展示，用于辅助识别主要分布特征和异常点。"
            f"{code_note}"
        )

    if len(field_names) == 1:
        return (
            f"本图围绕“{question}”展开，重点观察 {field_names[0]} 这一指标，"
            "用于辅助识别主要分布特征和异常点。"
            f"{code_note}"
        )

    return (
        "本图在完成必要的数据清洗后，对相关指标进行可视化展示，用于辅助识别主要分布特征和异常点。"
        f"{code_note}"
    )


def _build_quant_analysis_section(
    document: Any,
    results: list[ExecutionResult],
    report_data: ReportData,
) -> None:
    """Render the quantitative analysis section with Heading 2 subsections."""
    _add_section_heading(
        document,
        ANALYSIS_SECTION_TITLE,
        level=1,
        start_on_new_page=True,
    )

    analysis_sections = report_data["section_2_analysis"]
    if len(analysis_sections) != len(results):
        raise ValueError("section_2_analysis length must match the number of task results.")

    for index, (section_item, result) in enumerate(zip(analysis_sections, results), start=1):
        sub_title = section_item["sub_title"].strip() or f"任务 {index}：{result['question_zh']}"
        _add_section_heading(
            document,
            sub_title,
            level=2,
            start_on_new_page=False,
        )
        _add_body_paragraphs(document, f"研究问题：{result['question_zh']}")
        _add_label_paragraph(document, "分析方法：")
        _add_body_paragraphs(document, _build_analysis_method_text(result))
        _add_label_paragraph(document, "可视化图表：")
        _add_figure(document, result["image_path"])
        _add_label_paragraph(document, "图表分析：")
        _add_body_paragraphs(document, section_item["content"] or result["analysis_text"])


def _rebuild_report_body(
    docx_path: Path,
    results: list[ExecutionResult],
    report_data: ReportData,
) -> None:
    """Rebuild the report body with standard Word headings for TOC support."""
    _require_docx_runtime()
    document = Document(str(docx_path))
    _configure_document_styles(document)
    _clear_template_body(document)
    _enable_update_fields_on_open(document)
    existing_toc_count = _normalize_existing_toc_fields(document)
    if existing_toc_count > 0:
        LOGGER.info("检测到模板已有 %s 个目录字段，复用并规范目录指令。", existing_toc_count)
    else:
        _insert_toc_page(document)

    _add_section_heading(
        document,
        report_data["section_1_intro"]["title"],
        level=1,
        start_on_new_page=False,
    )
    _add_body_paragraphs(document, report_data["section_1_intro"]["content"])

    _build_quant_analysis_section(document, results, report_data)

    _add_section_heading(
        document,
        report_data["section_3_mechanism"]["title"],
        level=1,
        start_on_new_page=True,
    )
    _add_body_paragraphs(document, report_data["section_3_mechanism"]["content"])

    for reflection_item in report_data["section_4_reflection"]:
        _add_section_heading(
            document,
            reflection_item["sub_title"],
            level=1,
            start_on_new_page=True,
        )
        _add_body_paragraphs(document, reflection_item["content"])

    document.save(str(docx_path))


def _get_auto_toc_config() -> tuple[bool, str]:
    """Return auto TOC settings, falling back to enabled Word COM refresh."""
    if get_settings is None:
        return True, "word_com"
    try:
        settings = get_settings()
    except Exception as exc:
        LOGGER.warning("读取自动目录配置失败，将使用默认 Word COM 后处理：%s", exc)
        return True, "word_com"
    return settings.enable_auto_toc, settings.auto_toc_backend


def refresh_docx_toc(file_path: str | Path) -> bool:
    """Refresh a DOCX table of contents using the configured backend."""
    try:
        enable_auto_toc, auto_toc_backend = _get_auto_toc_config()
        if not enable_auto_toc:
            LOGGER.warning("自动目录刷新已关闭，跳过 DOCX 目录后处理。")
            return False
        if auto_toc_backend == "none":
            LOGGER.warning("AUTO_TOC_BACKEND=none，跳过 DOCX 目录后处理。")
            return False

        LOGGER.info("准备使用 %s backend 刷新 DOCX 目录。", auto_toc_backend)
        if auto_toc_backend == "word_com":
            return _refresh_toc_with_word_com(file_path)
        if auto_toc_backend == "libreoffice":
            return _refresh_toc_with_libreoffice(file_path)
        if auto_toc_backend == "aspose":
            return _refresh_toc_with_aspose(file_path)

        LOGGER.warning("不支持的自动目录刷新后端：%s，跳过后处理。", auto_toc_backend)
        return False
    except Exception as exc:
        LOGGER.warning("DOCX 目录后处理失败，将保留当前文档并继续流程：%s", exc)
        return False


def _update_story_range_fields(doc: Any) -> None:
    """Update fields in all Word story ranges such as headers and footers."""
    try:
        story_ranges = doc.StoryRanges
    except Exception as exc:
        LOGGER.warning("无法读取 Word story ranges，跳过页眉页脚字段刷新：%s", exc)
        return

    for story_range in story_ranges:
        current_range = story_range
        while current_range is not None:
            try:
                field_count = int(current_range.Fields.Count)
                if field_count > 0:
                    current_range.Fields.Update()
            except Exception as exc:
                LOGGER.warning("刷新某个 Word story range 字段失败，继续处理其他区域：%s", exc)
            try:
                current_range = current_range.NextStoryRange
            except Exception:
                current_range = None


def _refresh_toc_with_libreoffice(file_path: str | Path) -> bool:
    """Reserved LibreOffice backend entry point."""
    LOGGER.warning("libreoffice backend 尚未实现，跳过 DOCX 目录后处理：%s", file_path)
    return False


def _refresh_toc_with_aspose(file_path: str | Path) -> bool:
    """Reserved Aspose backend entry point."""
    LOGGER.warning("aspose backend 尚未实现，跳过 DOCX 目录后处理：%s", file_path)
    return False


def _refresh_toc_with_word_com(file_path: str | Path) -> bool:
    """使用 Word COM 在后台静默刷新目录与全部域，成功返回 True。"""
    absolute_file_path = os.path.abspath(str(file_path))
    LOGGER.info("开始执行 Word COM 目录刷新后处理：%s", absolute_file_path)

    if not os.path.exists(absolute_file_path):
        LOGGER.warning("需要刷新的 Word 文档不存在：%s", absolute_file_path)
        return False

    if platform.system() != "Windows":
        LOGGER.warning("当前环境不是 Windows，跳过 Word COM 自动目录刷新。")
        return False

    try:
        import win32com.client as win32  # type: ignore[import-not-found]
        import pythoncom  # type: ignore[import-not-found]
    except ImportError:
        LOGGER.warning(
            "当前环境缺少 win32com.client，无法自动刷新 Word 目录。"
            "请先执行 `pip install pywin32`。"
        )
        return False

    word = None
    doc = None
    com_initialized = False

    try:
        try:
            pythoncom.CoInitialize()
            com_initialized = True
            word = win32.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            LOGGER.info("Word COM 启动成功，已设置为后台静默模式。")

            LOGGER.info("正在打开文档。")
            doc = word.Documents.Open(
                FileName=absolute_file_path,
                ConfirmConversions=False,
                ReadOnly=False,
                AddToRecentFiles=False,
            )

            toc_count = int(doc.TablesOfContents.Count)
            if toc_count > 0:
                LOGGER.info("检测到 %s 个目录对象，开始逐个刷新。", toc_count)
                for index in range(1, toc_count + 1):
                    LOGGER.info("正在刷新第 %s 个目录。", index)
                    doc.TablesOfContents(index).Update()
                LOGGER.info("TOC 更新成功。")
            else:
                LOGGER.info("文档中未检测到显式目录对象，继续执行全部域更新。")

            LOGGER.info("正在更新文档中的全部域。")
            doc.Fields.Update()
            LOGGER.info("全文字段更新成功。")

            LOGGER.info("正在更新页眉页脚和其他 story ranges 中的字段。")
            _update_story_range_fields(doc)

            LOGGER.info("正在保存刷新后的文档。")
            doc.Save()
            LOGGER.info("DOCX 保存成功。")
            LOGGER.info("Word 目录刷新完成。")
            return True
        except Exception as refresh_exc:
            LOGGER.warning("自动刷新 Word 目录失败，将保留当前文档并继续流程：%s", refresh_exc)
            return False
    finally:
        if doc is not None:
            try:
                LOGGER.info("正在关闭 Word 文档句柄。")
                doc.Close(SaveChanges=True)
            except Exception as close_exc:
                LOGGER.warning("关闭 Word 文档时出现异常：%s", close_exc)

        if word is not None:
            try:
                LOGGER.info("正在退出 Word 后台进程。")
                word.Quit()
            except Exception as quit_exc:
                LOGGER.warning("退出 Word 进程时出现异常：%s", quit_exc)
        if com_initialized:
            try:
                pythoncom.CoUninitialize()
            except Exception as com_exc:
                LOGGER.warning("释放 Word COM 上下文时出现异常：%s", com_exc)


def refresh_word_toc(file_path: str | Path) -> bool:
    """Compatibility wrapper for the legacy Word COM-only TOC refresh entry."""
    return _refresh_toc_with_word_com(file_path)


def _resolve_path(path_text: str) -> Path:
    """Resolve a file path relative to the project root when needed."""
    path = Path(path_text).expanduser()
    if path.is_absolute():
        return path.resolve()
    return (PROJECT_ROOT / path).resolve()


def _resolve_existing_path(path_text: str, label: str) -> Path:
    """Resolve a path and ensure it exists."""
    resolved_path = _resolve_path(path_text)
    if not resolved_path.exists():
        raise FileNotFoundError(f"{label} not found: {resolved_path}")
    return resolved_path


def _validate_result_item(item: dict[str, Any], index: int) -> ExecutionResult:
    """Validate a single renderer input item."""
    if not isinstance(item, dict):
        raise ValueError(f"results[{index}] must be a dictionary.")

    missing_keys = REQUIRED_RESULT_KEYS.difference(item.keys())
    if missing_keys:
        raise ValueError(f"results[{index}] is missing keys: {sorted(missing_keys)}")

    task_id = item.get("task_id")
    if not isinstance(task_id, int):
        raise ValueError(f"results[{index}]['task_id'] must be an integer.")

    string_fields = {
        "question_zh",
        "image_path",
        "analysis_text",
        "code_snippet",
        "prepare_code",
        "plot_code",
        "exploration_output",
        "cleaning_summary",
        "problem_solution",
        "reflection_hint",
    }
    normalized: ExecutionResult = {
        "task_id": task_id,
        "question_zh": "",
        "image_path": "",
        "analysis_text": "",
        "code_snippet": "",
        "prepare_code": "",
        "plot_code": "",
        "exploration_output": "",
        "cleaning_summary": "",
        "problem_solution": "",
        "reflection_hint": "",
        "column_mapping": {},
    }
    for field_name in string_fields:
        value = item.get(field_name)
        if not isinstance(value, str) or not value.strip():
            raise ValueError(f"results[{index}]['{field_name}'] must be a non-empty string.")
        normalized[field_name] = value.strip()  # type: ignore[literal-required]

    normalized["image_path"] = str(
        _resolve_existing_path(normalized["image_path"], "Image").resolve()
    )

    column_mapping = item.get("column_mapping")
    if not isinstance(column_mapping, dict):
        raise ValueError(f"results[{index}]['column_mapping'] must be a dictionary.")
    normalized["column_mapping"] = {
        str(key).strip(): str(value).strip()
        for key, value in column_mapping.items()
        if str(key).strip() and str(value).strip()
    }
    return normalized


def _normalize_results(results: list[dict[str, Any]]) -> list[ExecutionResult]:
    """Validate and normalize the full renderer input list."""
    if not isinstance(results, list) or not results:
        raise ValueError("results must be a non-empty list of execution result dictionaries.")
    normalized_results = [
        _validate_result_item(item=item, index=index)
        for index, item in enumerate(results)
    ]
    return sorted(normalized_results, key=lambda item: item["task_id"])


def _validate_report_data(report_data: dict[str, Any]) -> ReportData:
    """Validate the TOC-friendly report data produced by Node 3.5."""
    if not isinstance(report_data, dict):
        raise ValueError("report_data must be a dictionary.")

    report_keys = set(report_data.keys())
    if report_keys != REQUIRED_REPORT_KEYS:
        raise ValueError(
            "report_data must contain exactly these keys: "
            f"{sorted(REQUIRED_REPORT_KEYS)}"
        )

    def validate_section_block(value: Any, field_name: str) -> ReportSectionBlock:
        if not isinstance(value, dict) or set(value.keys()) != {"title", "content"}:
            raise ValueError(
                f"report_data['{field_name}'] must be an object containing 'title' and 'content'."
            )
        title = value.get("title")
        content = value.get("content")
        if not isinstance(title, str) or not title.strip():
            raise ValueError(f"report_data['{field_name}']['title'] must be a non-empty string.")
        if not isinstance(content, str) or not content.strip():
            raise ValueError(
                f"report_data['{field_name}']['content'] must be a non-empty string."
            )
        return {
            "title": title.strip(),
            "content": content.strip(),
        }

    def validate_section_item(value: Any, field_name: str, index: int) -> ReportSectionItem:
        if not isinstance(value, dict) or set(value.keys()) != {"sub_title", "content"}:
            raise ValueError(
                f"report_data['{field_name}'][{index}] must contain 'sub_title' and 'content'."
            )
        sub_title = value.get("sub_title")
        content = value.get("content")
        if not isinstance(sub_title, str) or not sub_title.strip():
            raise ValueError(
                f"report_data['{field_name}'][{index}]['sub_title'] must be a non-empty string."
            )
        if not isinstance(content, str) or not content.strip():
            raise ValueError(
                f"report_data['{field_name}'][{index}]['content'] must be a non-empty string."
            )
        return {
            "sub_title": sub_title.strip(),
            "content": content.strip(),
        }

    section_2_analysis_raw = report_data.get("section_2_analysis")
    if not isinstance(section_2_analysis_raw, list) or not section_2_analysis_raw:
        raise ValueError("report_data['section_2_analysis'] must be a non-empty list.")
    section_4_reflection_raw = report_data.get("section_4_reflection")
    if not isinstance(section_4_reflection_raw, list) or not section_4_reflection_raw:
        raise ValueError("report_data['section_4_reflection'] must be a non-empty list.")

    return {
        "section_1_intro": validate_section_block(report_data["section_1_intro"], "section_1_intro"),
        "section_2_analysis": [
            validate_section_item(item, "section_2_analysis", index)
            for index, item in enumerate(section_2_analysis_raw)
        ],
        "section_3_mechanism": validate_section_block(
            report_data["section_3_mechanism"],
            "section_3_mechanism",
        ),
        "section_4_reflection": [
            validate_section_item(item, "section_4_reflection", index)
            for index, item in enumerate(section_4_reflection_raw)
        ],
    }


def _validate_data_summary(data_summary: dict[str, Any]) -> DataSummary:
    """Validate the renderer's dataset summary payload."""
    if not isinstance(data_summary, dict):
        raise ValueError("data_summary must be a dictionary.")

    missing_keys = REQUIRED_DATA_SUMMARY_KEYS.difference(data_summary.keys())
    if missing_keys:
        raise ValueError(f"data_summary is missing keys: {sorted(missing_keys)}")

    normalized: DataSummary = {
        "dataset_path": "",
        "dataset_name": "",
        "file_type": "",
        "shape_text": "",
        "info_text": "",
        "missing_summary_text": "",
        "duplicate_count_text": "",
        "preview_text": "",
        "load_code": "",
    }
    for key in REQUIRED_DATA_SUMMARY_KEYS:
        value = data_summary.get(key)
        if not isinstance(value, str) or not value.strip():
            raise ValueError(f"data_summary['{key}'] must be a non-empty string.")
        normalized[key] = value.strip()  # type: ignore[literal-required]

    resolved_dataset_path = _resolve_existing_path(normalized["dataset_path"], "Dataset path")
    normalized["dataset_path"] = resolved_dataset_path.as_posix()
    return normalized


def _validate_report_title(report_title: str) -> str:
    """Validate and normalize the report title."""
    if not isinstance(report_title, str) or not report_title.strip():
        raise ValueError("report_title must be a non-empty string.")
    return report_title.strip()


def _find_reflection_content(
    reflection_items: list[ReportSectionItem],
    *,
    keywords: tuple[str, ...],
    field_name: str,
) -> str:
    """Pick one reflection block by subtitle keywords."""
    for item in reflection_items:
        if any(keyword in item["sub_title"] for keyword in keywords):
            return item["content"]
    raise ValueError(f"Unable to map report_text['section_4_reflection'] to {field_name}.")


def _compose_source_code(item: ExecutionResult) -> str:
    """Build a clean code block for report embedding."""
    code_parts = [
        part.strip()
        for part in (item["prepare_code"], item["plot_code"])
        if part.strip()
    ]
    if code_parts:
        return "\n\n".join(code_parts)
    return item["code_snippet"].strip()


def _build_template_results(
    template: Any,
    results: list[ExecutionResult],
    report_data: ReportData,
) -> list[TemplateLoopItem]:
    """Build the loop payload consumed by report_template.docx."""
    analysis_sections = report_data["section_2_analysis"]
    if len(analysis_sections) != len(results):
        raise ValueError("section_2_analysis length must match the number of task results.")

    rendered_items: list[TemplateLoopItem] = []
    for result, section_item in zip(results, analysis_sections):
        rendered_items.append(
            {
                "question_zh": result["question_zh"],
                "analysis_text": section_item["content"] or result["analysis_text"],
                "source_code": _compose_source_code(result),
                "image": InlineImage(template, result["image_path"], width=Mm(140)),
            }
        )
    return rendered_items


def _build_template_context(
    template: Any,
    results: list[ExecutionResult],
    report_data: ReportData,
    data_summary: DataSummary,
    report_title: str,
) -> dict[str, Any]:
    """Flatten structured renderer inputs into the final template context."""
    del data_summary  # Reserved for future template placeholders.
    return {
        "report_title": report_title,
        "abstract": report_data["section_1_intro"]["content"],
        "detailed_analysis": report_data["section_3_mechanism"]["content"],
        "problems_encountered": _find_reflection_content(
            report_data["section_4_reflection"],
            keywords=("问题", "解决"),
            field_name="problems_encountered",
        ),
        "final_thoughts": _find_reflection_content(
            report_data["section_4_reflection"],
            keywords=("思考", "总结"),
            field_name="final_thoughts",
        ),
        "results": _build_template_results(template, results, report_data),
    }


def render_report(
    results: list[dict[str, Any]],
    report_text: dict[str, Any],
    data_summary: dict[str, Any],
    report_title: str = "SmartAnalyst Economist Report",
    output_dir: str | Path | None = None,
) -> dict[str, str | None]:
    """Render the academic Word report from the final docx template."""
    LOGGER.info("Starting template-driven academic report rendering.")

    if DocxTemplate is None or InlineImage is None or Mm is None:
        raise RuntimeError(
            "docxtpl is not installed, so template-based Word rendering is unavailable."
        ) from DOCX_TEMPLATE_IMPORT_ERROR
    _require_docx_runtime()

    normalized_results = _normalize_results(results)
    normalized_report = _validate_report_data(report_text)
    normalized_summary = _validate_data_summary(data_summary)
    normalized_title = _validate_report_title(report_title)
    renderer_paths = _build_renderer_paths(output_dir)
    _ensure_outputs_dir(renderer_paths.output_dir)

    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Word template not found: {TEMPLATE_PATH}")

    template = DocxTemplate(str(TEMPLATE_PATH))
    context = _build_template_context(
        template,
        normalized_results,
        normalized_report,
        normalized_summary,
        normalized_title,
    )
    template.render(context, autoescape=True)
    template.save(str(renderer_paths.docx_path))
    _rebuild_report_body(renderer_paths.docx_path, normalized_results, normalized_report)
    refresh_docx_toc(renderer_paths.docx_path)
    LOGGER.info("Academic Word report saved to %s", renderer_paths.docx_path)

    pdf_path: str | None = None
    try:
        if convert is None:
            raise RuntimeError("docx2pdf is not installed.")
        convert(str(renderer_paths.docx_path), str(renderer_paths.pdf_path))
        pdf_path = str(renderer_paths.pdf_path.resolve())
        LOGGER.info("Academic PDF report saved to %s", renderer_paths.pdf_path)
    except Exception as exc:
        LOGGER.warning(
            "PDF conversion failed for %s and will be skipped. "
            "This usually means Microsoft Word/docx2pdf is unavailable or access is denied: %s",
            renderer_paths.docx_path,
            exc,
        )

    return {
        "docx_path": str(renderer_paths.docx_path.resolve()),
        "pdf_path": pdf_path,
    }


def _build_import_cell() -> str:
    """Create the notebook import cell source."""
    return "\n".join(
        [
            "import pandas as pd",
            "import matplotlib.pyplot as plt",
            "from pathlib import Path",
        ]
    )


def _build_loading_cell(data_summary: DataSummary) -> str:
    """Create the notebook loading-and-exploration cell."""
    return "\n".join(
        [
            data_summary["load_code"],
            "print(df.shape)",
            "print(df.info())",
        ]
    )


def _build_loading_outputs(data_summary: DataSummary) -> list[Any]:
    """Create notebook outputs for the loading-and-exploration cell."""
    output_text = "\n".join(
        [
            data_summary["shape_text"],
            data_summary["info_text"],
            "None",
        ]
    )
    return [
        _new_output(
            output_type="stream",
            name="stdout",
            text=f"{output_text}\n",
        )
    ]


def _aggregate_column_mapping(results: list[ExecutionResult]) -> dict[str, str]:
    """Merge rename maps across tasks while preserving first-seen order."""
    aggregated: dict[str, str] = {}
    for item in results:
        for source_name, target_name in item["column_mapping"].items():
            aggregated.setdefault(source_name, target_name)
    return aggregated


def _aggregate_cleaning_notes(results: list[ExecutionResult]) -> list[str]:
    """Collect unique cleaning notes from task outputs."""
    notes: list[str] = []
    for item in results:
        for candidate in (item["cleaning_summary"], item["problem_solution"]):
            if candidate not in notes:
                notes.append(candidate)
    return notes


def _build_cleaning_markdown(results: list[ExecutionResult]) -> str:
    """Build notebook markdown for the cleaning section."""
    notes = _aggregate_cleaning_notes(results)
    lines = ["## 字段重命名与数据清洗", ""]
    for note in notes:
        lines.append(f"- {note}")
    return "\n".join(lines)


def _build_prepare_outputs(results: list[ExecutionResult]) -> list[Any]:
    """Use the first task's captured exploration output as notebook evidence."""
    if not results:
        return []
    return [
        _new_output(
            output_type="stream",
            name="stdout",
            text=f"{results[0]['exploration_output']}\n",
        )
    ]


def render_notebook(
    results: list[dict[str, Any]],
    report_data: dict[str, Any],
    data_summary: dict[str, Any],
    report_title: str = "SmartAnalyst Economist Report",
    output_dir: str | Path | None = None,
) -> str:
    """Render the sample-style Jupyter notebook."""
    LOGGER.info("Starting notebook rendering.")
    if nbf is None:
        LOGGER.warning(
            "nbformat is not installed; falling back to direct JSON notebook rendering: %s",
            NBF_IMPORT_ERROR,
        )

    normalized_results = _normalize_results(results)
    normalized_report = _validate_report_data(report_data)
    normalized_summary = _validate_data_summary(data_summary)
    normalized_title = _validate_report_title(report_title)
    renderer_paths = _build_renderer_paths(output_dir)
    _ensure_outputs_dir(renderer_paths.output_dir)

    if len(normalized_report["section_2_analysis"]) != len(normalized_results):
        raise ValueError("section_2_analysis length must match the number of task results.")

    notebook = _new_notebook()
    notebook["metadata"] = {
        "kernelspec": {
            "display_name": "Python 3",
            "language": "python",
            "name": "python3",
        },
        "language_info": {
            "name": "python",
        },
        "smartanalyst": {
            "mode": "Academic Authenticity",
        },
    }

    title_cell = _new_markdown_cell(
        "\n".join(
            [
                f"# {normalized_title}",
                "",
                f"数据文件：{normalized_summary['dataset_name']}",
                "",
                "本 notebook 保留了数据读取、探索、字段规范化、清洗与绘图过程。",
            ]
        )
    )
    intro_cell = _new_markdown_cell(
        f"## {normalized_report['section_1_intro']['title']}\n\n"
        f"{normalized_report['section_1_intro']['content']}"
    )
    import_cell = _new_code_cell(_build_import_cell())
    loading_cell = _new_code_cell(_build_loading_cell(normalized_summary))
    loading_cell["outputs"] = _build_loading_outputs(normalized_summary)
    loading_cell["execution_count"] = 1

    prepare_markdown_cell = _new_markdown_cell(_build_cleaning_markdown(normalized_results))
    prepare_code_cell = _new_code_cell(normalized_results[0]["prepare_code"])
    prepare_code_cell["outputs"] = _build_prepare_outputs(normalized_results)
    prepare_code_cell["execution_count"] = 2

    notebook["cells"].extend(
        [
            title_cell,
            intro_cell,
            import_cell,
            loading_cell,
            prepare_markdown_cell,
            prepare_code_cell,
        ]
    )

    execution_count = 3
    for section_item, item in zip(normalized_report["section_2_analysis"], normalized_results):
        task_markdown = _new_markdown_cell(
            "\n".join(
                [
                    f"## {section_item['sub_title']}",
                    "",
                    f"研究问题：{item['question_zh']}",
                ]
            )
        )
        plot_cell = _new_code_cell(item["plot_code"])
        plot_cell["execution_count"] = execution_count
        execution_count += 1

        image_name = Path(item["image_path"]).name
        result_markdown = _new_markdown_cell(
            "\n".join(
                [
                    f"![Task {item['task_id']} Output]({image_name})",
                    "",
                    "**简要分析**",
                    "",
                    item["analysis_text"],
                ]
            )
        )
        notebook["cells"].extend([task_markdown, plot_cell, result_markdown])

    notebook["cells"].append(
        _new_markdown_cell(
            f"## {normalized_report['section_3_mechanism']['title']}\n\n"
            f"{normalized_report['section_3_mechanism']['content']}"
        )
    )
    for reflection_item in normalized_report["section_4_reflection"]:
        notebook["cells"].append(
            _new_markdown_cell(
                f"## {reflection_item['sub_title']}\n\n{reflection_item['content']}"
            )
        )

    with renderer_paths.notebook_path.open("w", encoding="utf-8", newline="\n") as notebook_file:
        _write_notebook(notebook, notebook_file)

    LOGGER.info("Notebook saved to %s", renderer_paths.notebook_path)
    return str(renderer_paths.notebook_path.resolve())


def render_data_summary(
    results: list[dict[str, Any]],
    data_summary: dict[str, Any],
    output_dir: str | Path | None = None,
) -> str:
    """Write a plain-text cleaning summary file."""
    LOGGER.info("Starting cleaning-summary text rendering.")
    normalized_results = _normalize_results(results)
    normalized_summary = _validate_data_summary(data_summary)
    renderer_paths = _build_renderer_paths(output_dir)
    _ensure_outputs_dir(renderer_paths.output_dir)

    column_mapping = _aggregate_column_mapping(normalized_results)
    lines = [
        "SmartAnalyst Cleaning Summary",
        f"Generated At: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"Dataset: {normalized_summary['dataset_name']}",
        f"Shape: {normalized_summary['shape_text']}",
        f"Duplicate Rows: {normalized_summary['duplicate_count_text']}",
        "",
        "Missing Value Summary:",
        normalized_summary["missing_summary_text"],
        "",
        "Column Renaming:",
    ]
    if column_mapping:
        for source_name, target_name in column_mapping.items():
            lines.append(f"- {source_name} -> {target_name}")
    else:
        lines.append("- No renaming was needed.")

    lines.append("")
    lines.append("Task-Level Cleaning Notes:")
    for item in normalized_results:
        lines.append(f"- Task {item['task_id']} ({item['question_zh']}): {item['cleaning_summary']}")
        lines.append(f"  Problem/Solution: {item['problem_solution']}")

    renderer_paths.cleaning_summary_path.write_text(
        "\n".join(lines).strip() + "\n",
        encoding="utf-8",
        newline="\n",
    )
    LOGGER.info("Cleaning summary saved to %s", renderer_paths.cleaning_summary_path)
    return str(renderer_paths.cleaning_summary_path.resolve())


def run(
    results: list[dict[str, Any]],
    report_text: dict[str, Any],
    data_summary: dict[str, Any],
    report_title: str = "SmartAnalyst Economist Report",
    output_dir: str | Path | None = None,
) -> RenderArtifacts:
    """Convenience entry point for the renderer node."""
    report_artifacts = render_report(results, report_text, data_summary, report_title, output_dir=output_dir)
    notebook_path = render_notebook(results, report_text, data_summary, report_title, output_dir=output_dir)
    cleaning_summary_path = render_data_summary(results, data_summary, output_dir=output_dir)
    return {
        "docx_path": report_artifacts["docx_path"],
        "pdf_path": report_artifacts["pdf_path"],
        "notebook_path": notebook_path,
        "cleaning_summary_path": cleaning_summary_path,
    }


def _ensure_dummy_image(image_path: Path) -> None:
    """Create a tiny PNG image for standalone renderer testing."""
    if image_path.exists():
        return
    image_path.parent.mkdir(parents=True, exist_ok=True)
    tiny_png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAASwAAADICAIAAADdvUsCAAAABmJLR0QA/wD/AP+gvaeTAAAB"
        "TUlEQVR4nO3TMQ0AAAgDoGv/0qMFDxIFfXpn5gBA1w4A4LkDAOC5AwDguQMA4LkDAOC5AwDguQ"
        "MA4LkDAOC5AwDguQMA4LkDAOC5AwDguQMA4LkDAOC5AwDguQMA4LkDAOC5AwDguQMA4LkDAOC5"
        "AwDguQMA4LkDAOC5AwDguQMA4LkDAOC5AwDguQMA4LkDAOC5AwDguQMA4LkDAOC5AwDguQMA4L"
        "kDAOC5AwDguQMA4LkDAOC5AwDguQMA4LkDAOC5AwDguQMA4LkDAOC5AwDguQMA4LkDAOC5AwDg"
        "uQMA4LkDAOC5AwDguQMA4LkDAOC5AwDguQMA4LkDAOC5A4D9B3H1ArQk5m4AAAAASUVORK5CYII="
    )
    image_path.write_bytes(tiny_png)


def _ensure_self_test_dataset() -> Path:
    """Return an existing data file or create a tiny CSV fixture for self-tests."""
    data_dir = PROJECT_ROOT / "data"
    data_dir.mkdir(parents=True, exist_ok=True)
    for candidate in sorted(data_dir.iterdir(), key=lambda path: path.name.lower()):
        if candidate.is_file() and candidate.suffix.lower() in {".csv", ".xlsx", ".xls"}:
            return candidate.resolve()

    fixture_path = data_dir / "renderer_self_test.csv"
    fixture_path.write_text(
        "age,wage,edu\n"
        "22,5200,bachelor\n"
        "28,6800,master\n"
        "35,9100,master\n",
        encoding="utf-8",
        newline="\n",
    )
    return fixture_path.resolve()


if __name__ == "__main__":
    renderer_paths = _build_renderer_paths()
    _ensure_outputs_dir(renderer_paths.output_dir)
    dummy_image_path = renderer_paths.output_dir / "task_1.png"
    _ensure_dummy_image(dummy_image_path)
    dataset_path = _ensure_self_test_dataset()

    mock_results = [
        {
            "task_id": 1,
            "question_zh": "工资与年龄之间是否存在变化关系？",
            "image_path": str(dummy_image_path.resolve()),
            "analysis_text": "样本结果显示年龄增长与工资提升之间存在正向变化关系。",
            "code_snippet": "print('demo')",
            "prepare_code": "\n".join(
                [
                    "column_rename_map = {'age': '年龄', 'wage': '工资'}",
                    "df_clean = df.rename(columns=column_rename_map).copy()",
                    "print(df.shape)",
                    "print(df.info())",
                ]
            ),
            "plot_code": "\n".join(
                [
                    "data_plot = df_clean[['年龄', '工资']].dropna().copy()",
                    "plt.figure(figsize=(8, 5))",
                    "plt.plot(data_plot['年龄'], data_plot['工资'])",
                    f"plt.savefig(r'{dummy_image_path.resolve()}', dpi=300)",
                    "plt.close()",
                ]
            ),
            "exploration_output": "(3, 3)\n<class 'pandas.core.frame.DataFrame'>\nNone",
            "cleaning_summary": "我先把 age 和 wage 改写成中文字段名，再保留分析需要的列。",
            "problem_solution": "我发现原始字段中存在英文缩写，因此先统一改成中文，避免后续图表标签不自然。",
            "reflection_hint": "工资水平随年龄上升，可能对应经验积累效应。",
            "column_mapping": {"age": "年龄", "wage": "工资"},
        }
    ]
    mock_report = {
        "section_1_intro": {
            "title": "一、引言与数据清洗说明",
            "content": "本文基于样本数据，对变量关系进行了初步说明，并简要交代了字段重命名与基础清洗过程。",
        },
        "section_2_analysis": [
            {
                "sub_title": "（一）年龄与工资变化关系分析",
                "content": "图表显示年龄增长与工资提升之间存在较为稳定的正向关系，说明经验积累可能带来收入提升。",
            }
        ],
        "section_3_mechanism": {
            "title": "三、经济机制分析",
            "content": "从经验积累与岗位匹配角度看，年龄增长可能对应更强的人力资本沉淀，从而影响工资水平。",
        },
        "section_4_reflection": [
            {
                "sub_title": "四、遇到的问题及解决方法",
                "content": "在处理字段时，我首先遇到的是英文列名不够直观的问题，因此先完成字段重命名；在绘图前又检查了空值与重复值，保证图形结果更稳定。",
            },
            {
                "sub_title": "五、总结与思考",
                "content": "通过本次分析，我更能理解变量之间的结构关系，也意识到课堂作业中的字段规范化和图表表达同样重要。",
            },
        ],
    }
    mock_summary = {
        "dataset_path": str(dataset_path),
        "dataset_name": dataset_path.name,
        "file_type": dataset_path.suffix.lower().lstrip("."),
        "shape_text": "(3, 3)",
        "info_text": "<class 'pandas.core.frame.DataFrame'>\nRangeIndex: 3 entries, 0 to 2",
        "missing_summary_text": "age: 0\nwage: 0\nedu: 0",
        "duplicate_count_text": "0",
        "preview_text": "age wage edu",
        "load_code": "from pathlib import Path\ndata_path = Path('demo.csv')\ndf = pd.read_csv(data_path)",
    }

    artifacts = run(mock_results, mock_report, mock_summary, "Renderer Self-Test")
    print(json.dumps(artifacts, ensure_ascii=False, indent=2))
